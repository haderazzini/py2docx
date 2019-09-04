'''
###############################################################################
#                Procura no diretorio todas as pastas e arquivos                #
###############################################################################
#   version     date        author              comments                      #
###############################################################################
#   0.0.1       2019-09-03  Hader Azzini         innitial                     #
###############################################################################
'''

import os
import shutil
from docx import Document
from docx.shared import RGBColor

class ReaderDirectory():

    #def __init__(self,**kwargs):

    def get_list_dirs_first_level(self,directory):

        directory = os.path.abspath(directory)

        list_dirs = []

        for dirs in  os.walk(directory):
            list_dirs.append((dirs[1]))

        return list_dirs[0]
    
    def copy_to_temporary(self,directory,avoid_folders = ['temporary_py2pdf','__pycache__'],avoid_files = ['documentation.docx'], avoid_extentions =['cpython-36','cpython-37']):
        
        last_folder = os.path.split(directory)[1]
        new_prefix = last_folder+'\\'+'temporary_py2pdf'

        for root, dirs, files in os.walk(directory):

            temp_root = root.replace(last_folder,new_prefix)
            
            if os.path.split(directory)[1] not in avoid_folders:
                safe_create_directory(temp_root)

            if (files != []) & (os.path.split(directory)[1] not in avoid_folders):
                for filename in files:
                    if (filename not in avoid_files) & (filename.split('.')[1] not in avoid_extentions):
                        copy_file(root,temp_root,filename)

        temp_folder = directory.replace(last_folder,new_prefix)

        return temp_folder
    
    def read_all_files(self,directory,temp_folder,title='Software Documentation',docx_name='documentation',**kwargs):
        
        document = Document()
        document.add_heading(title, 0)

        for root, dirs, files in os.walk(temp_folder):

            if files !=[]:
                for filename in files:
                    print('------------------------')
                    print(filename)

                    if filename != docx_name+'.docx':
                        document.add_heading(filename.replace('.txt','.py'), level=1)
                        
                        full_name = os.path.join(root,filename)

                        f=open(full_name, "r", encoding="utf8")
                        if f.mode == 'r':
                            contents = f.read()

                            location = '#location of file: '+full_name.split(temp_folder)[1].replace('.txt','.py')
                            #document.add_paragraph(location)
                            write_in_another_color(document,location,'purple')

                            if os.stat(full_name).st_size == 0:
                                write_in_another_color(document,'# This file is empty!','red')
                            else:
                                document.add_paragraph(contents)
                            
                        f.close()
        
        document.save(os.path.join(directory,docx_name+'.docx'))

    def py2pdf(self,directory,**kwargs):

        temp_folder = self.copy_to_temporary(directory)

        self.read_all_files(directory,temp_folder,**kwargs)

        shutil.rmtree(temp_folder, ignore_errors=True)

    

def safe_create_directory(directory):

    if not os.path.exists(directory):
        os.makedirs(directory)


def copy_file(root,temporary_directory,filename):

    pre_name = filename.split('.py')[0]
    old_name = os.path.join(root,filename)
    new_name = os.path.join(temporary_directory,pre_name+'.txt')

    shutil.copy(old_name,new_name)

def write_in_another_color(document,text,color_name):


    font_colors = {'purple':RGBColor(0x42, 0x24, 0xE9),
                    'red': RGBColor(255, 0, 0)}

    run = document.add_paragraph().add_run(text)
    font = run.font
    font.color.rgb = font_colors[color_name]
